import streamlit as st
import os
from pathlib import Path
from datetime import datetime
import markdown
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from market_agent.crew import MarketAgent

st.set_page_config(
    page_title="Market Agent AI",
    page_icon="🚀",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for marketing theme
st.markdown("""
    <style>
    /* Main theme colors */
    :root {
        --primary-color: #6366f1;
        --secondary-color: #8b5cf6;
        --accent-color: #ec4899;
        --background-gradient: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
    }
    
    /* Header styling */
    .main-header {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        padding: 2rem;
        border-radius: 10px;
        margin-bottom: 2rem;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
    }
    
    .main-header h1 {
        color: white;
        font-size: 3rem;
        font-weight: 700;
        margin: 0;
        text-shadow: 2px 2px 4px rgba(0, 0, 0, 0.2);
    }
    
    .main-header p {
        color: #e0e7ff;
        font-size: 1.2rem;
        margin-top: 0.5rem;
    }
    
    /* Card styling */
    .info-card {
        background: white;
        padding: 1.5rem;
        border-radius: 10px;
        box-shadow: 0 2px 4px rgba(0, 0, 0, 0.1);
        border-left: 4px solid #6366f1;
        margin-bottom: 1rem;
    }
    
    .agent-card {
        background: linear-gradient(135deg, #f5f7fa 0%, #c3cfe2 100%);
        padding: 1rem;
        border-radius: 8px;
        margin: 0.5rem 0;
        border-left: 3px solid #8b5cf6;
    }
    
    /* Button styling */
    .stButton>button {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        font-weight: 600;
        border: none;
        padding: 0.75rem 2rem;
        border-radius: 8px;
        font-size: 1.1rem;
        transition: all 0.3s ease;
    }
    
    .stButton>button:hover {
        transform: translateY(-2px);
        box-shadow: 0 4px 12px rgba(102, 126, 234, 0.4);
    }
    
    /* Progress styling */
    .stProgress > div > div {
        background: linear-gradient(90deg, #667eea 0%, #764ba2 100%);
    }
    
    /* Sidebar styling */
    .css-1d391kg {
        background: linear-gradient(180deg, #667eea 0%, #764ba2 100%);
    }
    
    /* Success message */
    .success-box {
        background: linear-gradient(135deg, #84fab0 0%, #8fd3f4 100%);
        padding: 1.5rem;
        border-radius: 10px;
        color: #065f46;
        font-weight: 600;
        margin: 1rem 0;
    }
    
    /* Metric cards */
    .metric-card {
        background: white;
        padding: 1rem;
        border-radius: 8px;
        text-align: center;
        box-shadow: 0 2px 4px rgba(0, 0, 0, 0.1);
    }
    
    .metric-value {
        font-size: 2rem;
        font-weight: 700;
        color: #6366f1;
    }
    
    .metric-label {
        color: #6b7280;
        font-size: 0.9rem;
        margin-top: 0.5rem;
    }
    
    /* API Key section */
    .api-key-info {
        background: #fef3c7;
        padding: 1rem;
        border-radius: 8px;
        border-left: 4px solid #f59e0b;
        margin: 1rem 0;
    }
    
    .api-key-success {
        background: #d1fae5;
        padding: 0.75rem;
        border-radius: 6px;
        border-left: 3px solid #10b981;
        margin: 0.5rem 0;
    }
    </style>
""", unsafe_allow_html=True)

def convert_markdown_to_docx(markdown_text, product_idea):
    """Convert markdown report to Word document with formatting"""
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Add title
    title = doc.add_heading('Market Research Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.color.rgb = RGBColor(99, 102, 241)
    
    # Add product idea
    product_para = doc.add_paragraph()
    product_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    product_run = product_para.add_run(f'Product: {product_idea}')
    product_run.font.size = Pt(12)
    product_run.font.color.rgb = RGBColor(107, 114, 128)
    
    # Add date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f'Generated: {datetime.now().strftime("%B %d, %Y")}')
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = RGBColor(156, 163, 175)
    
    doc.add_paragraph()  # Spacing
    
    # Parse markdown and add to document
    lines = markdown_text.split('\n')
    
    for line in lines:
        line = line.strip()
        
        if not line:
            continue
            
        # Headers
        if line.startswith('# '):
            heading = doc.add_heading(line[2:], level=1)
            heading.runs[0].font.color.rgb = RGBColor(99, 102, 241)
        elif line.startswith('## '):
            heading = doc.add_heading(line[3:], level=2)
            heading.runs[0].font.color.rgb = RGBColor(139, 92, 246)
        elif line.startswith('### '):
            heading = doc.add_heading(line[4:], level=3)
            heading.runs[0].font.color.rgb = RGBColor(139, 92, 246)
        
        # Bullet points
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        
        # Numbered lists
        elif line[0:2].replace('.', '').isdigit():
            doc.add_paragraph(line.split('. ', 1)[1] if '. ' in line else line, style='List Number')
        
        # Bold text (simple detection)
        elif '**' in line:
            para = doc.add_paragraph()
            parts = line.split('**')
            for i, part in enumerate(parts):
                run = para.add_run(part)
                if i % 2 == 1:  # Odd indices are bold
                    run.bold = True
        
        # Regular paragraph
        else:
            doc.add_paragraph(line)
    
    # Save to BytesIO
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    return doc_io

def main():
    # Initialize session state for API keys
    if 'google_api_key' not in st.session_state:
        st.session_state.google_api_key = os.getenv('GOOGLE_API_KEY', '')
    if 'serper_api_key' not in st.session_state:
        st.session_state.serper_api_key = os.getenv('SERPER_API_KEY', '')
    if 'api_keys_configured' not in st.session_state:
        st.session_state.api_keys_configured = False
    
    # Header
    st.markdown("""
        <div class="main-header">
            <h1>🚀 Market Agent AI</h1>
            <p>AI-Powered Market Research & Business Analysis</p>
        </div>
    """, unsafe_allow_html=True)
    
    # Sidebar
    with st.sidebar:
        st.image("https://img.icons8.com/fluency/96/000000/artificial-intelligence.png", width=80)
        st.title("About")
        st.markdown("""
        **Market Agent AI** uses five specialized AI agents to conduct comprehensive market research:
        
        🔍 **Market Research Specialist**  
        📊 **Competitive Intelligence Analyst**  
        👥 **Customer Insights Researcher**  
        🎯 **Product Strategy Advisor**  
        💼 **Business Analyst**
        """)
        
        st.divider()
        
        # API Key Configuration Section
        st.markdown("### 🔑 API Configuration")
        
        with st.expander("⚙️ Configure API Keys", expanded=not st.session_state.api_keys_configured):
            st.markdown("""
            Enter your API keys below or configure them in the `.env` file.
            """)
            
            # Google Gemini API Key
            google_key_input = st.text_input(
                "Google Gemini API Key",
                value=st.session_state.google_api_key if st.session_state.google_api_key else "",
                type="password",
                help="Get your key at: https://makersuite.google.com/app/apikey",
                key="google_key_input"
            )
            
            if google_key_input:
                st.session_state.google_api_key = google_key_input
            
            st.markdown("[Get Google Gemini API Key →](https://makersuite.google.com/app/apikey)")
            
            st.markdown("---")
            
            # SerperDev API Key
            serper_key_input = st.text_input(
                "SerperDev API Key",
                value=st.session_state.serper_api_key if st.session_state.serper_api_key else "",
                type="password",
                help="Get your key at: https://serper.dev/ (2,500 free searches)",
                key="serper_key_input"
            )
            
            if serper_key_input:
                st.session_state.serper_api_key = serper_key_input
            
            st.markdown("[Get SerperDev API Key →](https://serper.dev/)")
            
            st.markdown("---")
            
            # Save button
            if st.button("💾 Save API Keys", use_container_width=True):
                if st.session_state.google_api_key and st.session_state.serper_api_key:
                    # Set environment variables for current session
                    os.environ['GOOGLE_API_KEY'] = st.session_state.google_api_key
                    os.environ['SERPER_API_KEY'] = st.session_state.serper_api_key
                    st.session_state.api_keys_configured = True
                    st.success("✅ API keys configured successfully!")
                    st.info("💡 Keys are saved for this session only. Add them to `.env` file for permanent storage.")
                else:
                    st.error("⚠️ Please enter both API keys")
        
        # API Key Status
        if st.session_state.google_api_key and st.session_state.serper_api_key:
            st.success("✅ API Keys Configured")
        else:
            st.warning("⚠️ API Keys Required")
            st.info("👆 Click 'Configure API Keys' above to get started")
        
        st.divider()
        
        st.markdown("### 📋 What You'll Get")
        st.markdown("""
        - Market size analysis (TAM/SAM/SOM)
        - Competitive landscape mapping
        - Customer personas & pain points
        - MVP feature prioritization
        - Pricing strategy
        - Financial projections
        - Go/No-Go recommendation
        """)
        
        st.divider()
        st.markdown("Built with ❤️ by Krishna")
    
    # Main content
    
    # API Key Warning (if not configured)
    if not st.session_state.google_api_key or not st.session_state.serper_api_key:
        st.warning("⚠️ **API Keys Required**: Please configure your API keys in the sidebar to start analyzing. Click '⚙️ Configure API Keys' in the left sidebar.")
    
    col1, col2 = st.columns([2, 1])
    
    with col1:
        st.markdown("### 💡 Enter Your Product Idea")
        product_idea = st.text_area(
            "Describe your AI product idea in detail",
            height=150,
            placeholder="Example: An AI-powered tool that summarizes YouTube videos and posts summaries on social media platforms like LinkedIn, Instagram, Facebook, X, and WhatsApp",
            help="Be as specific as possible about your product's features, target audience, and value proposition"
        )
    
    with col2:
        st.markdown("### 📊 Analysis Scope")
        st.markdown("""
        <div class="info-card">
            <strong>Report Includes:</strong><br>
            ✓ 3000+ word analysis<br>
            ✓ Market sizing & trends<br>
            ✓ 5-7 competitor profiles<br>
            ✓ Customer personas<br>
            ✓ Product roadmap<br>
            ✓ Financial projections<br>
            ✓ Investment thesis
        </div>
        """, unsafe_allow_html=True)
    
    st.divider()
    
    # Agent showcase
    st.markdown("### 🤖 Meet Your AI Research Team")
    
    agent_cols = st.columns(5)
    agents = [
        ("🔍", "Market Research", "Analyzes market size & trends"),
        ("📊", "Competitive Intel", "Maps competitive landscape"),
        ("👥", "Customer Insights", "Develops personas"),
        ("🎯", "Product Strategy", "Designs MVP & roadmap"),
        ("💼", "Business Analysis", "Creates investment thesis")
    ]
    
    for col, (icon, name, desc) in zip(agent_cols, agents):
        with col:
            st.markdown(f"""
                <div class="agent-card">
                    <div style="font-size: 2rem; text-align: center;">{icon}</div>
                    <div style="font-weight: 600; text-align: center; margin-top: 0.5rem;">{name}</div>
                    <div style="font-size: 0.8rem; text-align: center; color: #6b7280; margin-top: 0.3rem;">{desc}</div>
                </div>
            """, unsafe_allow_html=True)
    
    st.divider()
    
    # Action buttons
    col1, col2, col3 = st.columns([1, 1, 2])
    
    with col1:
        analyze_button = st.button("🚀 Start Analysis", type="primary", use_container_width=True)
    
    with col2:
        if os.path.exists("reports/report.md"):
            with open("reports/report.md", "r", encoding="utf-8") as f:
                report_content = f.read()
            
            # Convert to Word
            doc_io = convert_markdown_to_docx(report_content, product_idea if product_idea else "Previous Analysis")
            
            st.download_button(
                label="📥 Download Report (Word)",
                data=doc_io,
                file_name=f"market_research_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
    
    # Analysis execution
    if analyze_button:
        if not product_idea:
            st.error("⚠️ Please enter a product idea to analyze")
        elif not st.session_state.google_api_key or not st.session_state.serper_api_key:
            st.error("⚠️ Please configure your API keys in the sidebar first")
            st.info("👈 Click 'Configure API Keys' in the sidebar to add your Google Gemini and SerperDev API keys")
        else:
            # Set environment variables for the analysis
            os.environ['GOOGLE_API_KEY'] = st.session_state.google_api_key
            os.environ['SERPER_API_KEY'] = st.session_state.serper_api_key
            
            st.markdown("""
                <div class="success-box">
                    🎯 Analysis started! Your AI research team is now working...
                </div>
            """, unsafe_allow_html=True)
            
            # Progress tracking
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            tasks = [
                "🔍 Conducting market research...",
                "📊 Analyzing competitors...",
                "👥 Researching customer insights...",
                "🎯 Developing product strategy...",
                "💼 Creating business analysis..."
            ]
            
            try:
                # Initialize crew
                inputs = {"product_idea": product_idea}
                
                # Run analysis with progress updates
                for i, task in enumerate(tasks):
                    status_text.markdown(f"**{task}**")
                    progress_bar.progress((i + 1) * 20)
                    
                    if i == 0:
                        # Start the actual crew execution
                        with st.spinner("AI agents are collaborating..."):
                            result = MarketAgent().crew().kickoff(inputs=inputs)
                
                progress_bar.progress(100)
                status_text.markdown("**✅ Analysis complete!**")
                
                st.success("🎉 Market research report generated successfully!")
                
                # Display report
                if os.path.exists("reports/report.md"):
                    st.divider()
                    st.markdown("### 📄 Generated Report")
                    
                    with open("reports/report.md", "r", encoding="utf-8") as f:
                        report_content = f.read()
                    
                    # Show report in expandable section
                    with st.expander("📖 View Full Report", expanded=True):
                        st.markdown(report_content)
                    
                    # Download button
                    col1, col2, col3 = st.columns([1, 1, 2])
                    with col1:
                        doc_io = convert_markdown_to_docx(report_content, product_idea)
                        st.download_button(
                            label="📥 Download as Word Document",
                            data=doc_io,
                            file_name=f"market_research_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )
                    
                    with col2:
                        st.download_button(
                            label="📥 Download as Markdown",
                            data=report_content,
                            file_name=f"market_research_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md",
                            mime="text/markdown",
                            use_container_width=True
                        )
                
            except Exception as e:
                st.error(f"❌ An error occurred: {str(e)}")
                st.info("💡 Make sure your API keys are configured in the .env file")
    
    # Show existing report if available
    elif os.path.exists("reports/report.md"):
        st.divider()
        st.markdown("### 📄 Previous Report Available")
        st.info("👆 Click 'Download Report (Word)' above to download the last generated report, or start a new analysis")
        
        with st.expander("📖 View Previous Report"):
            with open("reports/report.md", "r", encoding="utf-8") as f:
                st.markdown(f.read())

if __name__ == "__main__":
    main()
